# -*- coding: utf-8 -*-
"""
Builds the SchuLyf Student Management System academic report (.docx).

Reproduces the formatting of the validated reference report:
  - Times New Roman 12pt body, 1.5 line spacing, justified
  - Running header (project title) + footer ("Written and Presented by ..." + page number)
  - Roman-numeral front matter, arabic-numeral body (separate sections)
  - Auto-numbered Figures/Tables (SEQ fields) + Table of Contents / List of Figures /
    List of Tables (TOC fields, refreshed by Word on open)

EDIT THE IDENTITY BLOCK BELOW (university, author, supervisor ...) then re-run:
    python build_report.py
"""

import re
import os
import struct
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --------------------------------------------------------------------------------------
# IDENTITY BLOCK — fill these in (placeholders shown in << >> are meant to be replaced)
# --------------------------------------------------------------------------------------
UNIVERSITY   = "<< UNIVERSITY NAME >>"
INSTITUTE    = "<< UNIVERSITY INSTITUTE / FACULTY >>"
ADDRESS      = "P.O. Box 0000 - Cameroon   |   Tel: (237) 000 00 00 00   |   www.example.com"
SCHOOL       = "School of Bachelors"
DEPARTMENT   = "DEPARTMENT OF COMPUTER ENGINEERING"
OPTION       = "OPTION: SOFTWARE ENGINEERING"
TOPIC        = "DESIGN AND IMPLEMENTATION OF A STUDENT MANAGEMENT SYSTEM"
SUBTITLE     = "Case study: SchuLyf"
AUTHOR_LABEL = "Written and Presented by:"
AUTHORS      = ["<< AUTHOR / GROUP NAME >>"]
MATRICULE    = "Matricule: << 00000 >>"
SUPERVISOR_LABEL = "Under the Supervision of:"
SUPERVISOR   = "<< Mr./Dr. SUPERVISOR NAME >>"
SUPERVISOR_TITLE = "<< Supervisor title >>"
ACADEMIC_YEAR = "2025 - 2026 Academic Year"

# Header / footer text repeated on every content page
RUNNING_HEADER = "DESIGN AND IMPLEMENTATION OF A STUDENT MANAGEMENT SYSTEM"
FOOTER_BYLINE  = "Written and Presented by << AUTHOR / GROUP NAME >>"

DIAG = os.path.join(os.path.dirname(os.path.abspath(__file__)), "diagrams")
SHOTS = os.path.join(os.path.dirname(os.path.abspath(__file__)), "screenshots")
OUT  = os.path.join(os.path.dirname(os.path.abspath(__file__)), "build",
                    "SchuLyf-Student-Management-System-Report.docx")

EMERALD = RGBColor(0x06, 0x5F, 0x46)
BLACK   = RGBColor(0x00, 0x00, 0x00)

document = Document()

# ======================================================================================
# Low-level helpers
# ======================================================================================

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_field(paragraph, instr, default_text=""):
    """Insert a Word complex field (begin / instrText / separate / default / end)."""
    run = paragraph.add_run()
    r = run._r
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve')
    instrText.text = instr
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = default_text
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    for el in (begin, instrText, sep, t, end):
        r.append(el)
    return run

def set_page_number_format(section, fmt, start=None):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))

def enable_update_fields_on_open():
    settings = document.settings.element
    upd = OxmlElement('w:updateFields')
    upd.set(qn('w:val'), 'true')
    settings.append(upd)

def add_runs(paragraph, text, base_size=None):
    """Add runs to a paragraph, honouring **bold** markers."""
    for part in re.split(r'(\*\*.*?\*\*)', text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2]); run.bold = True
        else:
            run = paragraph.add_run(part)
        if base_size:
            run.font.size = Pt(base_size)

# ======================================================================================
# Styles
# ======================================================================================

def configure_styles():
    normal = document.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    normal.font.color.rgb = BLACK
    rpr = normal.element.get_or_add_rPr().get_or_add_rFonts()
    rpr.set(qn('w:ascii'), 'Times New Roman')
    rpr.set(qn('w:hAnsi'), 'Times New Roman')
    rpr.set(qn('w:cs'), 'Times New Roman')
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size in (('Heading 1', 16), ('Heading 2', 14), ('Heading 3', 13)):
        st = document.styles[name]
        st.font.name = 'Times New Roman'
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = BLACK
        st.font.italic = False
        rf = st.element.get_or_add_rPr().get_or_add_rFonts()
        rf.set(qn('w:ascii'), 'Times New Roman'); rf.set(qn('w:hAnsi'), 'Times New Roman')
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    cap = document.styles['Caption']
    cap.font.name = 'Times New Roman'
    cap.font.size = Pt(11)
    cap.font.bold = True
    cap.font.italic = False
    cap.font.color.rgb = BLACK
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(10)

# ======================================================================================
# Content helpers
# ======================================================================================

def h1(text):
    document.add_paragraph(text, style='Heading 1')

def h2(text):
    document.add_paragraph(text, style='Heading 2')

def h3(text):
    document.add_paragraph(text, style='Heading 3')

def para(text):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs(p, text)
    return p

def bullet(text):
    p = document.add_paragraph(style='List Bullet')
    add_runs(p, text)
    return p

def numbered(text):
    p = document.add_paragraph(style='List Number')
    add_runs(p, text)
    return p

def page_break():
    document.add_page_break()

# Printable area of an A4 body page (8.27 x 11.69 in) after the 1.18"/1.0"
# side margins and 1.0" top/bottom margins, leaving headroom for the running
# header, the footer, and the figure caption. Any figure is scaled to fit
# INSIDE this box on BOTH axes, so a tall diagram can never overflow the page.
PAGE_MAX_W = 6.0
PAGE_MAX_H = 8.2


def _png_size(path):
    """(width, height) in pixels from a PNG's IHDR — no Pillow dependency."""
    with open(path, 'rb') as fh:
        head = fh.read(26)
    width, height = struct.unpack('>II', head[16:24])
    return width, height


def _fit(path, max_w, max_h):
    """Largest (w, h) in inches that fits max_w x max_h preserving aspect ratio."""
    px_w, px_h = _png_size(path)
    ratio = px_h / px_w
    w = max_w
    h = w * ratio
    if h > max_h:
        h = max_h
        w = h / ratio
    return w, h


def figure(filename, caption, width=PAGE_MAX_W, max_h=PAGE_MAX_H, src=DIAG):
    """Place an image scaled to fit the page on both axes.

    `width` is the maximum width (defaults to the full content width); the
    height is capped at `max_h` so tall diagrams shrink to fit instead of
    spilling past the bottom margin. `src` selects the diagrams or screenshots
    directory.
    """
    path = os.path.join(src, filename)
    w, h = _fit(path, width, max_h)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(w), height=Inches(h))
    add_caption('Figure', caption)


def figure_group(filenames, caption, width=PAGE_MAX_W, max_h=PAGE_MAX_H, src=DIAG):
    """Stack several images vertically under a single auto-numbered caption.

    Used when one logical figure is rendered as more than one panel (e.g. a
    use-case diagram split into two readable bands). The per-panel height cap
    keeps the whole stack inside the printable area.
    """
    per_h = (max_h - 0.15 * (len(filenames) - 1)) / len(filenames)
    for filename in filenames:
        path = os.path.join(src, filename)
        w, h = _fit(path, width, per_h)
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        run.add_picture(path, width=Inches(w), height=Inches(h))
    add_caption('Figure', caption)

def add_caption(label, text):
    p = document.add_paragraph(style='Caption')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{label} ")
    add_field(p, f' SEQ {label} \\* ARABIC ', default_text="1")
    p.add_run(f": {text}")

def placeholder(desc, caption, height_in=2.6):
    """A bordered single-cell box standing in for a screenshot."""
    tbl = document.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_table_borders(tbl, color="9CA3AF", sz=8)
    set_cell_background(cell, "F3F4F6")
    cell.width = Inches(6.0)
    # vertical padding via empty lines
    p0 = cell.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(height_in * 18)
    r = p0.add_run(f"[ SCREENSHOT PLACEHOLDER ]")
    r.bold = True; r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    p1 = cell.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(height_in * 18)
    r1 = p1.add_run(desc)
    r1.italic = True; r1.font.color.rgb = RGBColor(0x6B, 0x72, 0x80); r1.font.size = Pt(11)
    add_caption('Figure', caption)

def set_table_borders(table, color="000000", sz=6):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)
    tblPr.append(borders)

def make_table(caption, headers, rows, widths=None):
    add_caption('Table', caption)
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="000000", sz=6)
    hdr = table.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True; run.font.size = Pt(11)
        set_cell_background(hdr[i], "E7E7E7")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cp = cells[i].paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            cp.paragraph_format.space_after = Pt(2)
            add_runs(cp, str(val))
            for r in cp.runs:
                r.font.size = Pt(11)
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    document.add_paragraph()
    return table

# ======================================================================================
# Cover page
# ======================================================================================

def center_line(text, size=12, bold=False, italic=False, color=BLACK, space_after=2, caps=False):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text.upper() if caps else text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size); run.font.color.rgb = color
    return p

def boxed_topic(text, subtitle):
    tbl = document.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl, color="000000", sz=12)
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("TOPIC: " + text)
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = BLACK
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run(subtitle)
    r2.italic = True; r2.font.size = Pt(13)

def build_cover():
    center_line(UNIVERSITY, size=15, bold=True, space_after=2, caps=True)
    center_line(INSTITUTE, size=13, bold=True, space_after=4, caps=True)
    center_line(ADDRESS, size=9, space_after=10)
    center_line(SCHOOL, size=12, bold=True, space_after=2)
    center_line(DEPARTMENT, size=13, bold=True, space_after=16, caps=True)

    # crest / logo placeholder
    lp = document.add_paragraph()
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_after = Pt(16)
    lr = lp.add_run("[ INSTITUTION LOGO ]")
    lr.italic = True; lr.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF); lr.font.size = Pt(11)

    boxed_topic(TOPIC, SUBTITLE)
    document.add_paragraph()
    center_line(OPTION, size=13, bold=True, space_after=18)

    center_line(AUTHOR_LABEL, size=12, bold=True, space_after=4)
    for a in AUTHORS:
        center_line(a, size=13, bold=True, space_after=2, caps=True)
    center_line(MATRICULE, size=11, space_after=18)

    center_line(SUPERVISOR_LABEL, size=12, bold=True, space_after=4)
    center_line(SUPERVISOR, size=13, bold=True, space_after=2, caps=True)
    center_line(SUPERVISOR_TITLE, size=11, italic=True, space_after=22)

    center_line(ACADEMIC_YEAR, size=12, bold=True, space_after=2)

# ======================================================================================
# Sections / header / footer
# ======================================================================================

def style_header_footer(section):
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run(RUNNING_HEADER)
    hr.bold = True; hr.font.size = Pt(11); hr.font.name = 'Times New Roman'
    hr.font.color.rgb = BLACK
    # thin bottom rule under the running header
    pPr = hp._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '000000')
    pbdr.append(bottom); pPr.append(pbdr)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # tab stops: left byline, right page number
    fr = fp.add_run(FOOTER_BYLINE + "\t\t")
    fr.font.size = Pt(9); fr.font.name = 'Times New Roman'
    fr.italic = True
    pgrun_p = fp
    add_field(pgrun_p, ' PAGE ', default_text="1")
    for r in fp.runs:
        if r.font.size is None:
            r.font.size = Pt(9)

# ======================================================================================
# Build
# ======================================================================================

def build():
    configure_styles()

    # Section 0 — cover (no header/footer)
    sec0 = document.sections[0]
    sec0.left_margin = Inches(1.18)
    sec0.right_margin = Inches(1.0)
    sec0.top_margin = Inches(0.9)
    sec0.bottom_margin = Inches(0.9)
    build_cover()

    # Section 1 — front matter (roman numerals)
    document.add_section(WD_SECTION.NEW_PAGE)
    sec1 = document.sections[1]
    sec1.left_margin = Inches(1.18); sec1.right_margin = Inches(1.0)
    sec1.top_margin = Inches(1.0); sec1.bottom_margin = Inches(1.0)
    style_header_footer(sec1)
    set_page_number_format(sec1, 'lowerRoman', start=2)
    build_front_matter()

    # Section 2 — body (arabic numerals restart at 1)
    document.add_section(WD_SECTION.NEW_PAGE)
    sec2 = document.sections[2]
    sec2.left_margin = Inches(1.18); sec2.right_margin = Inches(1.0)
    sec2.top_margin = Inches(1.0); sec2.bottom_margin = Inches(1.0)
    style_header_footer(sec2)
    set_page_number_format(sec2, 'decimal', start=1)
    build_body()

    enable_update_fields_on_open()
    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    document.save(OUT)
    print("Saved:", OUT)

# ----- front matter --------------------------------------------------------------------

def toc_field(title, instr):
    h1(title)
    p = document.add_paragraph()
    add_field(p, instr,
              default_text="(Right-click and choose “Update Field” to populate, "
                           "or press Ctrl+A then F9 in Word.)")
    page_break()

def build_front_matter():
    toc_field("LIST OF TABLES", ' TOC \\h \\z \\c "Table" ')
    toc_field("LIST OF FIGURES", ' TOC \\h \\z \\c "Figure" ')
    toc_field("TABLE OF CONTENTS", ' TOC \\o "1-3" \\h \\z \\u ')

# ----- body ----------------------------------------------------------------------------

def build_body():
    general_introduction()
    structure_of_work()
    chapter_one()
    chapter_two()
    chapter_three()
    general_conclusion()
    references()

# Body content lives in report_content.py-style functions below.

# === GENERAL INTRODUCTION =============================================================

def general_introduction():
    h1("GENERAL INTRODUCTION")
    para(
        "The rapid evolution of Information and Communication Technologies (ICT) has "
        "transformed nearly every sector of human activity, and higher education is no "
        "exception. Universities the world over are replacing paper registers, hand-carried "
        "documents, and word-of-mouth communication with integrated digital platforms that "
        "manage the entire student lifecycle — from admission, through tuition payment and "
        "course work, to the publication of examination results. Such platforms are known "
        "generically as **Student Management Systems** (SMS) or Student Information Systems."
    )
    para(
        "In many Cameroonian institutions, however, several core administrative processes "
        "remain manual, fragmented, and error-prone. Admission is conducted in person: a "
        "candidate fills a physical application form, attaches supporting documents, and "
        "deposits the folder at the student-affairs office, then waits to be notified of a "
        "decision by mail. Tuition payment is validated by carrying a bank-deposit slip to "
        "the accounts office, where it is exchanged for a paper **school receipt** that is "
        "easily lost, forged, or shared. Access to examination halls during the exam period "
        "depends on a student's payment standing, which is computed by hand. Notices of a "
        "lecturer's absence travel informally and are often delivered late or distrusted. "
        "These weaknesses are a recurring source of friction, wasted time, and disputes "
        "between students and administrators."
    )
    para(
        "This project addresses those weaknesses through the design and implementation of "
        "**SchuLyf**, a web-based Student Management System tailored to the realities of a "
        "Cameroonian university. SchuLyf digitises five interlocking problem areas: online "
        "**admissions**, tamper-proof **payments and receipts**, **exam-gating and tuition "
        "deferrals**, **course management**, and **notifications** — all underpinned by "
        "role-based access control and an immutable audit trail that records every "
        "significant action."
    )
    para(
        "The system was developed using the **Laravel 13** PHP framework with an "
        "**Inertia.js + Vue 3** single-page front end, a **MySQL** database, and modern "
        "software-engineering practices including UML-based modelling, an action-oriented "
        "service layer, automated testing with Pest, and a security posture hardened against "
        "a documented audit. The remainder of this report presents the state of the art, the "
        "analysis and design of the system, and the results obtained."
    )
    page_break()

def structure_of_work():
    h1("STRUCTURE OF THE WORK IN PARTS")
    para("This work is organised into three chapters, the second of which is divided into "
         "three parts. The structure is as follows:")
    bullet("**Chapter 1 — State of the Art of Student Management Systems:** the history and "
           "evolution of student management, a survey of existing systems, and the "
           "conceptualization of the proposed system.")
    bullet("**Chapter 2 — System Analysis and Design**, divided into three parts:")
    p = document.add_paragraph(style='List Bullet 2')
    add_runs(p, "**Part One — Preliminary Study:** study of the existing (manual) system, "
                "its limitations, the proposed solutions, and the justification of the chosen "
                "solution.")
    p = document.add_paragraph(style='List Bullet 2')
    add_runs(p, "**Part Two — System Requirements:** the functional and non-functional "
                "requirements of the new system.")
    p = document.add_paragraph(style='List Bullet 2')
    add_runs(p, "**Part Three — System Design:** the architecture and UML models, the "
                "database design, and the needed resources and cost estimation.")
    bullet("**Chapter 3 — Results and Discussion:** the application interface, a discussion "
           "of how the implemented features meet the objectives, and the testing performed.")
    para("The work closes with a general conclusion and a list of references.")
    page_break()

# === CHAPTER 1 ========================================================================

def chapter_one():
    h1("CHAPTER 1: STATE OF THE ART OF STUDENT MANAGEMENT SYSTEMS")

    h2("1.1 Introduction")
    para(
        "Education is one of the most information-intensive activities of any society. A "
        "single university must keep accurate, current records of who has applied, who has "
        "been admitted, who has paid, who is enrolled in which course, who attended which "
        "session, and what marks each student earned. Historically these records were kept on "
        "paper and managed by hand. The growth of student populations and the maturity of ICT "
        "have made **Student Management Systems** indispensable: they automate admissions, "
        "registration, fee management, course administration, results processing, and "
        "communication, while improving accuracy, transparency, and accountability."
    )
    para(
        "This chapter reviews the state of the art of student management systems. It traces "
        "their historical evolution, surveys representative existing solutions, identifies "
        "their strengths and gaps, and sets out the conceptual foundations of the system "
        "proposed in this work."
    )

    h2("1.2 Historical Background and Evolution of Student Management Systems")
    para(
        "Record-keeping in education is as old as formal education itself. For centuries, "
        "student records — enrolment lists, fee ledgers, and examination results — were "
        "written by hand in registers and stored in filing cabinets. The methods, tools, and "
        "limitations evolved alongside the broader history of computing, as summarised in "
        "Table 1."
    )
    make_table(
        "Historical evolution of student management systems",
        ["Period", "Era", "Key characteristics", "Tools & methods", "Limitations"],
        [
            ["Before 1960", "Manual records", "Hand-written registers and fee ledgers; in-person admission", "Paper registers, filing cabinets", "Data loss, no search, slow retrieval"],
            ["1960s–1980s", "Mainframe era", "Centralised batch processing of student records", "Punch cards, mainframes", "High cost, no real-time access, specialist staff"],
            ["1980s–1990s", "Desktop databases", "Departmental record-keeping on personal computers", "dBASE, FoxPro, spreadsheets", "Data silos, single-machine, duplication"],
            ["1990s–2000s", "Client–server SIS", "Networked student-information systems on campus LANs", "Relational databases, LAN clients", "Limited remote access, costly licensing"],
            ["2000s–2010s", "Web-based systems", "Online portals and learning-management systems", "Web SIS, Moodle, student portals", "Integration cost, security concerns"],
            ["2010s–present", "Cloud & mobile", "Integrated cloud platforms with analytics and APIs", "Cloud SIS, web frameworks, mobile apps", "Subscription cost, data-sovereignty concerns"],
            ["Proposed", "SchuLyf", "Secure, role-based, tamper-proof receipts, immutable audit", "Laravel, Inertia/Vue, MySQL", "Requires hosting and internet access"],
        ],
        widths=[0.9, 1.0, 1.6, 1.4, 1.5],
    )

    h2("1.3 Existing Student Management Systems")
    para(
        "A number of student management and student-information systems are in use today, "
        "ranging from open-source learning platforms to enterprise-grade commercial suites. "
        "Each digitises part of the student lifecycle, but few are well adapted to the "
        "specific administrative realities — and budgets — of a single faculty in a "
        "developing-country university. Common limitations include high licensing cost, "
        "excessive complexity, a focus on the learning side at the expense of admissions and "
        "finance, and the absence of locally relevant features such as bank-slip validation "
        "and tamper-proof school receipts. Table 2 compares representative systems."
    )
    make_table(
        "Overview of existing student management systems",
        ["System", "Type / focus", "Strengths", "Limitations"],
        [
            ["Moodle", "Learning management system", "Open-source, rich course tools", "Not a full SIS; no admissions or fee-receipt workflow"],
            ["Fedena", "School/college ERP", "All-in-one modules", "Generic, heavyweight, paid advanced modules"],
            ["openSIS", "Open-source SIS", "Free, student-information focus", "K-12 oriented, dated interface, limited finance"],
            ["Ellucian Banner", "Enterprise SIS", "Comprehensive, scalable", "Very costly, complex, over-featured for a faculty"],
            ["PowerSchool", "Commercial SIS", "Mature, well supported", "Subscription cost, limited local customization"],
            ["Local manual system", "Paper-based", "Low cost, familiar to staff", "Data loss, forgeable receipts, no audit trail"],
            ["Proposed (SchuLyf)", "Custom web SMS", "Tailored, role-based, tamper-proof receipts, audited", "Requires internet and hosting"],
        ],
        widths=[1.3, 1.4, 1.9, 1.9],
    )

    h2("1.4 Conceptualization of the Proposed System")
    para(
        "The conceptualization of SchuLyf translates the day-to-day operations of student "
        "administration — admission, payment validation, exam clearance, course delivery, and "
        "communication — into a single, secure, role-based digital platform. The design is "
        "organised around the following principles:"
    )
    h3("a. Role-Based Access Control")
    para(
        "Six roles — **Applicant, Student, Lecturer, Accountant, Student Affairs Officer "
        "(SAO), and Administrator** — each see only the functions appropriate to their "
        "responsibilities. This protects sensitive operations (admission decisions, payment "
        "validation, results publication, user management) and keeps each user's workspace "
        "focused and uncluttered."
    )
    h3("b. Digitised Admissions")
    para(
        "Applicants apply online against cascading programme and level menus, upload the "
        "exact documents their programme requires, and track the outcome. An SAO triages and "
        "decides each application; on admission, the applicant atomically becomes a student "
        "with a uniquely generated **matricule**, replacing the manual paper-folder process."
    )
    h3("c. Tamper-Proof Payments and Receipts")
    para(
        "A student reports a bank deposit by uploading the slip, so a lost paper copy is no "
        "longer fatal. An accountant validates the deposit, and the system issues a single "
        "**HMAC-signed school receipt** that is bound to the student's identity and can be "
        "verified by anyone at a public endpoint — making forgery and reuse detectable."
    )
    h3("d. Exam-Gating and Deferrals")
    para(
        "Tuition is paid in dated installments. The system computes each student's "
        "**payment standing** live (cleared, blocked, or deferred) and uses it to gate "
        "examination and facility access, while an accountant-granted **deferral** can lift "
        "the gate for a defined period."
    )
    h3("e. Course Management and Communication")
    para(
        "Courses are planned and approved, sessions are scheduled, attendance and assignments "
        "are recorded, and continuous-assessment and examination marks are published with a "
        "student dispute channel. Lecturer-absence and session changes are pushed to the "
        "affected cohort through reliable email and in-app **notifications**."
    )
    h3("f. Accountability through an Immutable Audit Trail")
    para(
        "Every significant action — a decision, a validation, a role change, a login — is "
        "written to an append-only **audit log** that can never be edited or deleted through "
        "the application, giving administrators a tamper-evident record for monitoring and "
        "dispute resolution."
    )
    para(
        "Together these principles deliver a secure, scalable, and user-friendly platform "
        "that overcomes the limitations of manual student administration. This conceptual "
        "design is the foundation for the system architecture, database design, UML models, "
        "and user interfaces presented in Chapter 2."
    )
    page_break()

# === CHAPTER 2 ========================================================================

def chapter_two():
    h1("CHAPTER 2: SYSTEM ANALYSIS AND DESIGN")
    para(
        "This chapter presents the analysis and design of the SchuLyf Student Management "
        "System. It is divided into three parts: a preliminary study of the existing system "
        "and the justification of the chosen solution; the functional and non-functional "
        "requirements; and the system design, comprising the architecture, UML models, "
        "database design, and resource and cost estimation."
    )

    # ---- PART ONE ----
    h1("PART ONE: PRELIMINARY STUDY")

    h2("2.1 Study of the Existing System")
    para(
        "The first step in developing the system was to study how student administration is "
        "currently carried out in the modelled university. The processes are predominantly "
        "manual and paper-based:"
    )
    bullet("**Admission** — a candidate collects and fills a physical application form, "
           "attaches photocopies of credentials, and deposits the folder at the "
           "student-affairs office; a decision is reached manually and communicated by mail.")
    bullet("**Payment validation** — a student deposits tuition into one of the institute's "
           "bank accounts (UBA, AFG, AFRICLAND, SCB, SGC), then carries the bank slip to the "
           "accounts office where it is checked and exchanged for a paper school receipt.")
    bullet("**Exam clearance** — before examinations, staff manually verify that each student "
           "has paid the installment due; those short of the threshold are denied access "
           "unless granted a written deferral.")
    bullet("**Course information** — class schedules, lecturer-absence notices, attendance, "
           "and results are managed on paper or by word of mouth.")

    h3("2.1.2 Limitations of the Manual System")
    bullet("High risk of loss or damage to physical application folders and receipts.")
    bullet("A paper school receipt can be **forged, altered, or shared** between students.")
    bullet("Students who lose a bank slip cannot prove payment, causing disputes with staff.")
    bullet("Manual computation of payment standing is slow and inconsistent.")
    bullet("Absence notices and results travel informally, are delayed, and are distrusted.")
    bullet("No reliable, tamper-evident record of who did what and when (no audit trail).")
    para(
        "These limitations show that the manual system is inefficient, insecure, and "
        "difficult to scale, motivating a tailored digital solution."
    )

    h3("2.1.3 Proposed Solutions")
    para("Three approaches were considered to address the limitations identified above.")
    h3("Solution 1: Improved Manual System")
    para("Standardised forms, pre-printed serial-numbered receipts, and spreadsheet ledgers "
         "could reduce some errors. **Advantages:** low cost, familiar to staff. "
         "**Limitations:** still prone to loss and forgery, no real-time standing, no audit, "
         "and poor scalability.")
    h3("Solution 2: Adopt a Commercial / Off-the-Shelf SIS")
    para("A commercial student-information system or ERP could be purchased. "
         "**Advantages:** ready-made, feature-rich. **Limitations:** recurring licensing "
         "cost, excessive complexity, weak fit to local workflows such as bank-slip "
         "validation, and dependence on a third-party vendor.")
    h3("Solution 3: Custom Student Management System (the chosen solution)")
    para("A custom web application — SchuLyf — built specifically for the faculty's "
         "workflows, integrating role-based access, online admissions, tamper-proof "
         "receipts, exam-gating, course management, notifications, and an immutable audit "
         "log. **Advantages:** tailored, cost-effective in the long term, secure, and "
         "extensible. **Limitations:** requires initial development effort and ongoing "
         "maintenance.")
    make_table(
        "Comparison of the proposed solutions",
        ["Solution", "Key features", "Advantages", "Limitations"],
        [
            ["Improved manual", "Standard forms, serial receipts, spreadsheets", "Low cost, familiar", "Forgery and loss persist, no audit, poor scaling"],
            ["Commercial SIS", "Ready-made modules, broad feature set", "Fast to deploy, feature-rich", "Costly, complex, weak local fit, vendor lock-in"],
            ["Custom SMS (SchuLyf)", "Role-based access, HMAC receipts, exam-gating, audit", "Tailored, secure, cost-effective, extensible", "Development time and maintenance"],
        ],
        widths=[1.3, 1.9, 1.8, 1.5],
    )

    h3("2.1.4 Justification of the Chosen System")
    para("The custom system was chosen for the following reasons:")
    bullet("**Tailored functionality** — it implements exactly the faculty's processes "
           "(bank-slip validation, year-scoped matricules, installment-based exam-gating) "
           "rather than forcing them into a generic product.")
    bullet("**Cost-effectiveness** — it is built on a free, open-source stack and avoids "
           "recurring per-seat licensing fees.")
    bullet("**Security and accountability** — role-based access, hashed credentials, "
           "HMAC-signed receipts, and an immutable audit log directly target the trust and "
           "fraud problems of the manual process.")
    bullet("**Scalability and maintainability** — a clean, layered architecture lets new "
           "features be added without redesigning the system.")
    bullet("**Improved communication and data accuracy** — automated notifications and "
           "centralised, validated data reduce delays, duplication, and disputes.")

    # ---- PART TWO ----
    h1("PART TWO: SYSTEM REQUIREMENTS")
    h2("2.2 System Requirements")
    para(
        "System requirements define what the SchuLyf Student Management System must do "
        "(functional requirements) and the quality attributes under which it must operate "
        "(non-functional requirements). They were derived from the study of the existing "
        "system, stakeholder needs, and the project objectives."
    )

    h3("2.2.1 Functional Requirements")
    para("The functional requirements describe the core operations the system performs, "
         "grouped by module. Table 3 summarises the six user roles and their principal "
         "abilities, after which each module's requirements are detailed.")
    make_table(
        "User roles and principal abilities",
        ["Role", "Principal abilities"],
        [
            ["Applicant", "Register, submit and track an admission application, upload documents"],
            ["Student", "Report payments, view standing, request deferrals, view courses/attendance/assignments/results, raise disputes, print receipts"],
            ["Lecturer", "Manage assigned courses and plans, schedule sessions, mark attendance, set and grade assignments, record marks"],
            ["Accountant", "Validate/reject payments, issue receipts, review deferral requests, look up a student's standing"],
            ["Student Affairs Officer", "Triage and decide applications, manage courses and lecturers, approve plans, publish results, resolve disputes"],
            ["Administrator", "Manage users and reference data, configure fee schedules, view the audit log; may perform every other role's actions"],
        ],
        widths=[1.6, 4.9],
    )

    h3("A. Authentication, Authorization and User Management")
    bullet("Users authenticate with a single login field that accepts an **email, employee "
           "ID, phone number, or matricule**, resolved in one query.")
    bullet("Passwords are stored using strong one-way **hashing**; optional two-factor "
           "authentication (TOTP) and email verification are supported.")
    bullet("Access to every feature is restricted by **role**; sensitive actions are further "
           "protected by ability gates and per-resource ownership checks.")
    bullet("Administrators create staff accounts that receive a single-use **invite link** "
           "to set their own password; accounts can be deactivated, restored, and have their "
           "role changed.")
    bullet("Login attempts and password resets are **rate-limited** to resist brute force.")

    h3("B. Admissions")
    bullet("Applicants build an application against **cascading dropdowns** (degree "
           "programme → department → offering → level) and upload the documents required for "
           "that programme and level, in addition to the always-required national ID and "
           "birth certificate.")
    bullet("The system enforces **one open application at a time** per applicant.")
    bullet("An SAO **triages** an application through reversible interim states and records "
           "a single terminal **decision** (admitted, rejected, or waitlisted).")
    bullet("On admission, the applicant is **promoted to a student** atomically: a student "
           "profile and a unique year-scoped matricule are created and the Student role is "
           "assigned.")
    bullet("A returning student's prior enrolment can be **restored** instead of issuing a "
           "fresh record.")

    h3("C. Payments and Tamper-Proof Receipts")
    bullet("A student **reports a bank deposit** by selecting the bank, entering the amount "
           "and reference, and uploading the slip; the document is stored so a lost paper "
           "copy is no longer fatal.")
    bullet("An accountant **validates or rejects** each submission; validation **issues an "
           "immutable, HMAC-signed school receipt** bound to the student's identity.")
    bullet("Anyone may **verify a receipt** at a public endpoint; a forged number, altered "
           "amount, or reuse by another student is detected and reported as invalid.")
    bullet("Administrators configure **fee schedules** and dated **installments** per "
           "programme, level, and academic year.")

    h3("D. Exam-Gating and Tuition Deferrals")
    bullet("The system computes each student's **payment standing** live — cleared, blocked, "
           "or deferred — from the fee schedule, validated payments, and active deferrals.")
    bullet("A blocked student may **request a deferral**; an accountant approves it with an "
           "extended deadline or rejects it with a reason.")
    bullet("Gate-facing staff can **look up any student's standing by matricule**.")

    h3("E. Course Management")
    bullet("An SAO publishes courses and **assigns a lecturer**; the lecturer drafts a "
           "**course plan** for SAO approval.")
    bullet("On an approved plan the lecturer **schedules sessions**, **marks attendance** "
           "for the cohort, and creates and **grades assignments**.")
    bullet("Continuous-assessment and examination marks are recorded, then **published** by "
           "an SAO; the final score and letter grade are computed automatically.")
    bullet("Students view **published results** and may **raise a dispute**, which an SAO "
           "resolves.")

    h3("F. Notifications")
    bullet("The system sends **email and in-app notifications** for admission decisions, "
           "payment outcomes, deferral decisions, and course-session cancellations or "
           "reschedules.")
    bullet("Session-change notifications are fanned out to the **entire affected cohort**, "
           "replacing unreliable word-of-mouth notices.")

    h3("G. Audit and Reporting")
    bullet("Every significant write is recorded in an **immutable audit log** with the "
           "actor, action, before/after changes, and request context.")
    bullet("Administrators read the audit log through a filterable modal; role and "
           "application dashboards summarise counts and recent activity.")

    h3("2.2.2 Non-Functional Requirements")
    para("The non-functional requirements define the quality attributes and constraints of "
         "the system, summarised in Table 4 and detailed below.")
    make_table(
        "Non-functional requirements",
        ["Attribute", "Requirement"],
        [
            ["Security", "Hashed passwords, role-based access, ability gates, HMAC-signed receipts, hardened file viewers, rate limiting, input validation"],
            ["Performance", "Responsive pages, indexed queries, cached reference data, queued email so requests stay fast"],
            ["Reliability", "Atomic database transactions, terminal-state re-guards, graceful error handling"],
            ["Usability", "Clean, consistent role-based interface; clear forms and feedback toasts"],
            ["Scalability", "Layered architecture; new modules added without redesign; cache and queue back-ends swappable"],
            ["Maintainability", "Action-class pattern, typed routes, automated tests, and documentation"],
            ["Compatibility", "Runs in modern browsers; responsive layout for desktop and mobile"],
            ["Data integrity", "Foreign-key constraints, soft deletes, immutable receipt and audit records"],
        ],
        widths=[1.5, 5.0],
    )
    h3("A. Security")
    para("Authentication is backed by Laravel Fortify on a session guard. Passwords use "
         "strong hashing; in production a 12-character policy with breach checking is "
         "enforced. Authorization is layered — role middleware, ability gates, and "
         "per-resource ownership checks. School receipts are HMAC-signed and the inline file "
         "viewers apply a strict MIME allowlist and sandboxing content-security policy. "
         "Sensitive endpoints are rate-limited.")
    h3("B. Performance")
    para("Reference data (departments, offerings, document types) is served through a "
         "read-through cache; database queries are backed by composite indexes; and email "
         "and notifications are dispatched on a queue after the database transaction commits, "
         "so user requests are never blocked by slow I/O.")
    h3("C. Reliability and Data Integrity")
    para("Every state change runs inside a database transaction with a row lock and a "
         "terminal-state re-guard, so a record cannot be decided twice and partial writes "
         "roll back cleanly. Receipts and audit records are immutable; most other tables use "
         "soft deletes so data is recoverable.")
    h3("D. Usability, Scalability and Maintainability")
    para("The interface is built with the PrimeVue component library on a consistent emerald "
         "design system, presenting each role with only the actions relevant to it. The "
         "layered, action-oriented architecture and a comprehensive automated test suite make "
         "the system straightforward to extend and maintain.")

    # ---- PART THREE ----
    h1("PART THREE: SYSTEM DESIGN, RESOURCES AND COST ESTIMATION")

    h2("2.3 System Architecture")
    para(
        "SchuLyf follows a layered, server-driven single-page-application architecture. The "
        "Vue 3 front end communicates with the Laravel back end through Inertia.js: a page "
        "navigation is an XHR that returns a JSON page object which the client renders "
        "without a full reload. On the server, routes pass through an authentication, "
        "email-verification, and role-middleware stack before reaching a thin controller. "
        "**Every state change is encapsulated in a single-purpose Action class** that opens "
        "a transaction, locks the target row, re-guards its status, writes the change, "
        "records an audit entry, and dispatches any notification after the transaction "
        "commits. Figure 1 shows the layered architecture."
    )
    figure("10-architecture.png", "Layered system architecture of SchuLyf", width=5.6)
    para(
        "The technology stack is summarised in Table 5. The choice of a mature open-source "
        "stack keeps licensing cost at zero while providing strong security defaults, a "
        "powerful object-relational mapper, and a modern, reactive user interface."
    )
    make_table(
        "Technology stack",
        ["Layer", "Technology", "Role"],
        [
            ["Language / framework", "PHP 8.4, Laravel 13", "Application framework and domain logic"],
            ["Authentication", "Laravel Fortify", "Login, registration, reset, email verification, 2FA"],
            ["SPA bridge", "Inertia.js v3", "Server-routed single-page application"],
            ["Front end", "Vue 3, TypeScript, Vite", "Reactive interface components"],
            ["UI library", "PrimeVue (Aura), Tailwind CSS v4", "Forms, tables, dialogs, design system"],
            ["Database", "MySQL", "Relational persistence"],
            ["Cache / queue", "File or Redis cache; database queue", "Reference cache, queued mail and notifications"],
            ["Testing", "Pest v4, Laravel Pint", "Automated tests and code formatting"],
        ],
        widths=[1.5, 1.9, 3.1],
    )

    h2("2.4 System Design (UML Diagrams)")
    para("The system was modelled with the Unified Modeling Language (UML) before and during "
         "implementation. This section presents the use-case, class, sequence, activity, and "
         "state diagrams, together with the entity-relationship model of the database.")

    h3("2.4.1 Use Case Diagram")
    para("The use-case diagram (Figure 2) presents the major actors — Applicant, Student, "
         "Lecturer, Accountant, Student Affairs Officer, Administrator, and an anonymous "
         "public verifier — and the principal functions each performs within the system. "
         "For readability it is shown in two panels: the upper panel covers admissions, "
         "payments and exam clearance, and the lower panel covers course management and "
         "administration.")
    figure_group(
        ["01-usecase-1.png", "01-usecase-2.png"],
        "Use case diagram of the Student Management System",
    )

    h3("2.4.2 Class Diagram")
    para("The class diagram (Figure 3) shows the main domain entities, their key attributes "
         "and operations, and the relationships between them — for example, a validated "
         "payment submission yields exactly one immutable school receipt, and an admitted "
         "application gives rise to a student profile.")
    figure("02-class.png", "Class diagram of the principal domain entities", width=6.3)

    h3("2.4.3 Sequence Diagrams")
    para("Sequence diagrams illustrate the interactions between objects over time for the "
         "system's most important workflows.")
    para("Figure 4 shows the **admission decision**: when an SAO admits an application, the "
         "decision action re-fetches the record under a lock, updates its status, mints a "
         "student profile and matricule, assigns the Student role, writes an audit entry, "
         "and — only after the transaction commits — dispatches the decision email.")
    figure("05-seq-decision.png", "Sequence diagram — SAO admission decision", width=6.0)
    para("Figure 5 shows **payment validation and receipt issuance**: the accountant's "
         "action locks the submission, re-guards its terminal status, marks it validated, "
         "draws the next receipt number from a locked per-year counter, computes the HMAC "
         "signature over the receipt number, matricule, amount, and academic year, and "
         "creates the immutable receipt.")
    figure("06-seq-payment.png", "Sequence diagram — payment validation and receipt issuance", width=6.0)
    para("Figure 6 shows **public receipt verification**: the verifier supplies only a "
         "receipt number, the system recomputes the HMAC from the receipt's currently bound "
         "identity, and reveals the bound details only when the signature matches; an "
         "unknown, forged, or reused receipt reads simply as invalid.")
    figure("07-seq-verify.png", "Sequence diagram — public receipt verification (HMAC)", width=5.4)

    h3("2.4.4 Activity Diagrams")
    para("Activity diagrams describe the flow of control through a process. Figure 7 models "
         "the **online application** funnel from registration to submission, including the "
         "one-open-application guard.")
    figure("04-activity-application.png", "Activity diagram — online application submission", width=5.0)
    para("Figure 8 models the **payment-standing computation** that drives exam-gating: the "
         "service compares the installments due so far against validated payments and, on a "
         "shortfall, checks for an active deferral before returning cleared, deferred, or "
         "blocked.")
    figure("08-activity-standing.png", "Activity diagram — payment standing and exam-gating", width=5.4)

    h3("2.4.5 State Diagram")
    para("Figure 9 shows the **application lifecycle** as a state machine: an application "
         "moves through reversible interim states under SAO triage before reaching one of "
         "the terminal decisions, of which only the admitted and withdrawn paths affect a "
         "student record.")
    figure("09-state-application.png", "State diagram — application lifecycle", width=5.8)

    h3("2.4.6 Database Design (Entity-Relationship Model)")
    para("The database comprises twenty-four domain tables plus supporting framework and "
         "counter tables. Status and type columns are stored as strings backed by PHP "
         "enumerations rather than native database enumerations, and almost every table "
         "carries soft-delete and audit capabilities. Figure 10 shows the "
         "entity-relationship diagram of the core tables, and Table 6 groups the principal "
         "entities by bounded context.")
    figure("03-er.png", "Entity-relationship diagram of the core database", width=6.3)
    make_table(
        "Principal database entities by bounded context",
        ["Context", "Tables"],
        [
            ["Identity & roles", "users, roles, role_user, student_profiles, lecturer_profiles, accountant_profiles, sao_profiles"],
            ["Admissions", "departments, program_offerings, document_types, level_credential_requirements, applications, application_documents"],
            ["Payments & receipts", "fee_schedules, fee_installments, payment_submissions, school_receipts"],
            ["Exam-gating", "tuition_deferrals (payment standing is derived, never stored)"],
            ["Courses", "courses, course_sessions, attendance_records, assignments, assignment_submissions, course_results, result_disputes"],
            ["Notifications & audit", "notifications, audit_logs"],
        ],
        widths=[1.7, 4.8],
    )

    h2("2.5 Needed Resources and Cost Estimation")
    para("The development of the system required human, hardware, and software resources. "
         "Because the project was built by a small student team on an open-source stack, the "
         "dominant cost is human effort; most software was free of charge.")

    h3("2.5.1 Needed Resources")
    para("**Human resources** — a system analyst to define requirements, a UI/UX designer "
         "to model the diagrams and interfaces, front-end and back-end developers to build "
         "the application, a tester to verify functionality and security, and a project "
         "manager to coordinate the work.")
    para("**Hardware resources** — a development laptop, an external drive for backups, and "
         "an internet connection for research and deployment, as detailed in Table 8.")
    para("**Software resources** — the Laravel framework, the Inertia/Vue front-end stack, "
         "MySQL, Visual Studio Code, a UML modelling tool, and Git for version control — all "
         "open-source or free.")

    h3("2.5.2 Cost Estimation")
    para("The estimated cost of the project, in CFA francs (XAF), is broken down into human, "
         "hardware, and software resources in Tables 7 to 10.")
    make_table(
        "Cost of human resources",
        ["Designation", "Rate/day (XAF)", "Days", "Total (XAF)"],
        [
            ["Project manager", "3 000", "7", "21 000"],
            ["System analyst / designer", "3 000", "10", "30 000"],
            ["UI/UX designer", "2 500", "8", "20 000"],
            ["Front-end developer", "5 500", "15", "82 500"],
            ["Back-end developer", "6 500", "18", "117 000"],
            ["Tester", "3 500", "10", "35 000"],
            ["Total 1", "", "", "305 500"],
        ],
        widths=[2.4, 1.5, 1.0, 1.6],
    )
    make_table(
        "Cost of hardware resources",
        ["Designation", "Unit price (XAF)", "Quantity", "Total (XAF)"],
        [
            ["Development laptop", "400 000", "1", "400 000"],
            ["External hard drive (1 TB)", "35 000", "1", "35 000"],
            ["Modem + internet (project duration)", "30 000", "1", "30 000"],
            ["Total 2", "", "", "465 000"],
        ],
        widths=[2.7, 1.6, 1.1, 1.6],
    )
    make_table(
        "Cost of software resources",
        ["Designation", "Price (XAF)", "Notes"],
        [
            ["Laravel, Inertia, Vue, MySQL", "0", "Open-source"],
            ["Visual Studio Code, Git", "0", "Free"],
            ["UML modelling tool", "0", "Free / community edition"],
            ["Domain name + hosting (1 year)", "25 000", "Optional, for deployment"],
            ["Total 3", "", "25 000"],
        ],
        widths=[2.7, 1.4, 2.4],
    )
    make_table(
        "Total cost estimation",
        ["Item", "Amount (XAF)"],
        [
            ["Total 1 — human resources", "305 500"],
            ["Total 2 — hardware resources", "465 000"],
            ["Total 3 — software resources", "25 000"],
            ["Grand total", "795 500"],
        ],
        widths=[3.5, 2.0],
    )
    page_break()

# === CHAPTER 3 ========================================================================

def chapter_three():
    h1("CHAPTER 3: RESULTS AND DISCUSSION")
    para(
        "This chapter presents the practical outcomes of the SchuLyf Student Management "
        "System. It shows the principal application interfaces, discusses how each "
        "implemented feature meets the objectives defined in the previous chapters, and "
        "summarises the testing carried out to validate the system. The screenshots below "
        "were captured from the running application."
    )

    h2("3.1 Application Interface Screenshots")

    h3("3.1.1 Login Page")
    para("The login page authenticates every type of user through a single field that "
         "accepts an email, employee ID, phone number, or matricule, and redirects each user "
         "to the dashboard for their role.")
    figure("01-login.png", "Login page", src=SHOTS)

    h3("3.1.2 Applicant — Application Form")
    para("The application form guides an applicant through cascading programme and level "
         "menus and presents exactly the document slots required for the chosen programme, "
         "in addition to the always-required national ID and birth certificate.")
    figure("02-applicant-application-form.png", "Applicant application form", src=SHOTS)

    h3("3.1.3 Student Affairs Officer — Review Queue")
    para("The SAO review screen lists applications in a filterable, paginated queue and "
         "allows the officer to triage an application through interim states and record a "
         "final admission decision.")
    figure("03-sao-review-queue.png", "SAO application review queue", src=SHOTS)

    h3("3.1.4 Accountant — Payment Review")
    para("The accountant reviews a reported deposit, views the uploaded bank slip inline, "
         "and validates it — issuing a tamper-proof school receipt — or rejects it with a "
         "reason.")
    figure("04-accountant-payment-review.png",
           "Accountant payment review with the inline bank-slip viewer", src=SHOTS)

    h3("3.1.5 Student — Payments and Printable Receipt")
    para("The student payments screen shows the fee schedule, the validated-paid total, and "
         "the live payment standing, alongside a printable, QR-bearing school receipt that "
         "links to the public verification page.")
    figure("05-student-payments.png", "Student payments screen and payment standing", src=SHOTS)
    figure("06-student-receipt.png", "Printable, QR-bearing school receipt", src=SHOTS)

    h3("3.1.6 Public Receipt Verification")
    para("Anyone can verify a receipt number at a public page; an authentic receipt reveals "
         "the bound student identity and amount, while a forged or reused receipt is reported "
         "as invalid.")
    figure("07-receipt-verify.png", "Public receipt verification page", src=SHOTS)

    h3("3.1.7 Administrator — Dashboard and Audit Log")
    para("The administrator dashboard summarises users by role and applications by status "
         "and opens the immutable audit-log modal, where every significant action can be "
         "filtered by actor, action, subject, and date.")
    figure("08-admin-dashboard.png", "Administrator dashboard", src=SHOTS)
    figure("09-admin-audit-log.png", "Immutable audit-log modal with filters", src=SHOTS)

    h2("3.2 Testing and Validation")
    para(
        "The system was validated with an extensive automated test suite written with the "
        "Pest framework, comprising several hundred feature and unit tests in addition to "
        "browser smoke tests. The tests cover authentication and authorization for every "
        "role, the admission decision and student-promotion flow, payment validation and "
        "receipt signature verification, the payment-standing and deferral logic, the full "
        "course-management lifecycle, and the immutability of the audit log and receipts. "
        "Code style is enforced automatically with Laravel Pint, and the front end is "
        "type-checked and linted. This automated coverage gives confidence that the "
        "implemented behaviour matches the specified requirements and that future changes "
        "will not silently break existing functionality."
    )

    h2("3.3 Discussion")
    para(
        "The implemented system meets the objectives set out in Chapter 1. Online admissions "
        "replace the paper-folder process and produce a unique matricule automatically. The "
        "tamper-proof, HMAC-signed receipt with public verification resolves the twin "
        "problems of lost and forged paper receipts that motivated the project. Live payment "
        "standing and accountant-granted deferrals make exam-gating consistent and "
        "transparent. Course management brings planning, attendance, assignments, results, "
        "and disputes into one auditable workflow, and reliable cohort notifications replace "
        "distrusted word-of-mouth absence notices. Above all, the immutable audit log gives "
        "the institution a tamper-evident record of every significant action, directly "
        "addressing the accountability gaps of the manual system."
    )
    page_break()

# === CONCLUSION & REFERENCES ==========================================================

def general_conclusion():
    h1("GENERAL CONCLUSION")
    para(
        "The Student Management System developed in this project successfully addresses the "
        "major challenges of the traditional, manual student-administration processes of a "
        "Cameroonian university — processes that are fragmented, paper-bound, and prone to "
        "loss, forgery, and dispute. By digitising and automating admissions, payment "
        "validation, exam-gating, course management, and communication, SchuLyf significantly "
        "improves efficiency, transparency, and accountability."
    )
    para(
        "The application provides a secure, role-based platform for six categories of user. "
        "Applicants apply and track decisions online; students report payments, monitor their "
        "standing, follow their courses, and verify their receipts; lecturers run their "
        "courses end to end; accountants validate payments and issue tamper-proof receipts; "
        "student-affairs officers decide admissions and oversee the academic workflow; and "
        "administrators manage the platform and audit it. The use of a robust relational "
        "database, HMAC-signed receipts, and an immutable audit log ensures reliable, "
        "tamper-evident record-keeping throughout."
    )
    para(
        "Overall, the project achieved its objectives by delivering a functional, scalable, "
        "and well-tested Student Management System tailored to local needs. It demonstrates "
        "how a carefully designed, security-conscious web application can transform student "
        "administration. Future work could extend the platform with online payment-gateway "
        "integration, a dedicated mobile application, timetable generation, and richer "
        "analytics dashboards — all of which the layered architecture is designed to "
        "accommodate without redesign."
    )
    page_break()

def references():
    h1("REFERENCES")
    refs = [
        "Laravel. “Laravel 13 Documentation.” https://laravel.com/docs (accessed 2026).",
        "Inertia.js. “The Modern Monolith.” https://inertiajs.com (accessed 2026).",
        "Vue.js. “Vue 3 Documentation.” https://vuejs.org (accessed 2026).",
        "PrimeVue. “Vue UI Component Library.” https://primevue.org (accessed 2026).",
        "Tailwind CSS. “Documentation.” https://tailwindcss.com (accessed 2026).",
        "MySQL. “MySQL 8.0 Reference Manual.” https://dev.mysql.com/doc (accessed 2026).",
        "OWASP Foundation. “OWASP Top Ten.” https://owasp.org/www-project-top-ten (accessed 2026).",
        "H. Krawczyk, M. Bellare, R. Canetti. “HMAC: Keyed-Hashing for Message Authentication,” RFC 2104, IETF, 1997.",
        "Pest. “An Elegant PHP Testing Framework.” https://pestphp.com (accessed 2026).",
        "OMG. “Unified Modeling Language (UML) Specification, v2.5.1,” Object Management Group, 2017.",
    ]
    for i, r in enumerate(refs, 1):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        add_runs(p, f"[{i}]  {r}")

if __name__ == "__main__":
    build()
